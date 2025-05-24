import streamlit as st
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document as LangDocument
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import ChatOllama
import numpy as np
import faiss
from sklearn.feature_extraction.text import TfidfVectorizer
import tempfile
import os

# Load PDF
def load_pdf(file_path):
    loader = PyPDFLoader(file_path)
    return loader.load()

# Load DOCX
def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [LangDocument(page_content=full_text)]

# Load TXT
def load_txt(file_path):
    loader = TextLoader(file_path)
    return loader.load()

# Simple TF-IDF Embedding Wrapper
class TfidfEmbeddings:
    def __init__(self):
        self.vectorizer = TfidfVectorizer()

    def fit_transform(self, texts):
        return self.vectorizer.fit_transform(texts).toarray()

    def transform(self, query):
        return self.vectorizer.transform([query]).toarray()

# Streamlit UI Config
st.set_page_config(page_title="Offline PDF/DOC/TXT Q&A App", page_icon="📖")

st.title("📖 Offline PDF / DOCX / Text Q&A (Local Ollama + TF-IDF + FAISS)")

uploaded_files = st.file_uploader("Upload files (PDF, DOCX, or TXT)", accept_multiple_files=True, type=["pdf", "docx", "txt"])

if uploaded_files:
    docs = []
    for uploaded_file in uploaded_files:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.read())
            file_path = tmp_file.name

        ext = os.path.splitext(uploaded_file.name)[1].lower()
        if ext == ".pdf":
            docs.extend(load_pdf(file_path))
        elif ext == ".docx":
            docs.extend(load_docx(file_path))
        elif ext == ".txt":
            docs.extend(load_txt(file_path))

    if docs:
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        split_docs = text_splitter.split_documents(docs)

        # Extract plain text from LangDocument objects
        texts = [doc.page_content for doc in split_docs]

        # Initialize TF-IDF embeddings and fit
        embedding_model = TfidfEmbeddings()
        doc_vectors = embedding_model.fit_transform(texts)

        # FAISS Index setup
        dimension = doc_vectors.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(np.array(doc_vectors))

        # Ollama LLM
        llm = ChatOllama(model="mistral")  # swap 'tinyllama' with 'llama2' or 'mistral' if you prefer

        question = st.text_input("💡 Ask a question based on your uploaded files")

        if question:
            with st.spinner("Retrieving and answering..."):
                # Embed the query
                query_vec = embedding_model.transform(question)

                # Search top 3 relevant chunks
                D, I = index.search(query_vec, 3)
                retrieved_docs = [texts[i] for i in I[0]]

                # Concatenate context
                context = "\n".join(retrieved_docs)

                # Pass context and question to LLM
                prompt = f"Answer the question based on the following context:\n\n{context}\n\nQuestion: {question}"
                response = llm.invoke(prompt)

                st.success("✅ Answer:")
               
                st.write(response.content)


            st.markdown("---")
            if st.button("Show Processed Text Chunks"):
                for doc in split_docs:
                    st.write(doc.page_content)
    else:
        st.warning("No valid text found in uploaded files.")

else:
    st.info("Upload PDF, DOCX, or TXT files to get started.")
